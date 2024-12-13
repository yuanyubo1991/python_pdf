import PyPDF2
import docx

pdffilename = r"C:\Users\test.pdf"
wordfilename = r"C:\Users\test.docx"

def pdf_read():
    pdffile = PyPDF2.PdfReader(open(pdffilename, 'rb'))
    pdflines = len(pdffile.pages)
    pdfpagenum = int(input("please input the pdf file page number (1 - " + str(pdflines) + "):\n"))
    print(pdffilename)
    print("\tPages: " + str(pdflines))
    pdfpage = pdffile.pages[pdfpagenum - 1]
    pdfpagetext = pdfpage.extract_text()
    print(pdfpagetext)
    
def word_read():
    docfile = docx.Document(wordfilename)
    wordpara = len(docfile.paragraphs)
    print(wordfilename)
    print(wordpara)
    print(docfile.paragraphs[1].text)
    print(len(docfile.paragraphs[1].runs))
    print(docfile.paragraphs[1].runs[0].text)
    docfile.paragraphs[0].text = "good"
    docfile.paragraphs[1].text = "job"
    docfile.add_paragraph("finish!")
    docfile.save(wordfilename)

    
pdf_read()
word_read()
